#!/usr/bin/env python3
"""Generate a DOCX from a Slack thread export."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

SLACK_URL = "https://rippling.slack.com/archives/C0B4BLM7VDJ/p1781039024934769"
CHANNEL = "#proj-automated-amendment-submission"
TITLE = "Unsubmitted Amendments Backlog — Slack Thread"

PARENT = {
    "author": "Nicky DePaul",
    "email": "ndepaul@rippling.com",
    "time": "2026-06-09 17:03:44 EDT",
    "text": """Tony, James and I met today about how we're going to handle the backlog of unsubmitted amendments sitting in user's amendments tab. We have around 40K unsubmitted amendments right now! Our goal is to cut this number down to the minimum required to be filed. No dates on this yet, we are still working through the specifics.

Big picture takeaway: Ops will likely receive a major spike in submitted amendments regardless of how we handle this. We are choosing to accept near-term pain for long-term benefit.

Decisions made:
• No amendments generated prior to 2023 can be filed (approx 20k). We will need to figure out how to remove those from the queue before the automated submission begins
• We will give clients 30 days to submit their pending amendments or tell us they already filed outside of Rippling
• We will charge for all submitted amendments by default, as we presume any amendments related to incidents were handled at the time and are not sitting in this queue.

Next steps:
• Identify which amendments were user generated vs system generated (DRI: Abhishek Shah)
    ◦ We need to be able to explain why an amendment was generated, which has been a challenge in the past.
    ◦ If we cannot make this determination, we may not be able to charge for any of these backlogged amendments.
• Deepen the analysis of the existing backlog to understand % breakdown of amendments that should vs. should not be filed in order to cut down on the number of amendments that will pass to Ops (DRI: Tony Henriquez)
    ◦ System generated vs user-generated
    ◦ Churned vs. active clients
    ◦ Tax only vs. includes wages
    ◦ PEO vs. non-PEO
    ◦ marked filed externally vs. not
    ◦ W-2C(s) needed vs. not
• Calculate the total $ value of amended wages to serve as an input into legal decision making (DRI: Tony Henriquez)
• Draft GTM plan and customer comms (DRI: Nicky DePaul)
    ◦ Provide x days notice for clients to submit amendments manually before the automated process takes over
    ◦ Work with Legal to create a notice that admins must sign before the deadline. This would likely qualify as a payroll blocker at the end of the grace period—the admin must opt-out of us filing the amendment and assume responsibility/liability for filing, or opt-in to us filing their outstanding backlog of amendments.
    ◦ We would likely roll this out in batches, x% of clients per month, to keep Ops burden manageable
• Select a test cohort of 100 or so clients (removing any named accounts / high value etc) to understand the CS volume and patterns this might drive (DRI: Nicky DePaul)""",
}

REPLIES = [
    {
        "author": "Nicky DePaul",
        "email": "ndepaul@rippling.com",
        "time": "2026-06-09 17:04:42 EDT",
        "text": "cc Kushal Shah, Kitty Kwan, Krishna Mohan",
    },
    {
        "author": "Kitty Kwan",
        "email": "kitty@rippling.com",
        "time": "2026-06-09 18:53:33 EDT",
        "text": "Thanks - very helpful summary. Tony Henriquez do you also want the breakdown by agency so that we know which ones are 1) fully automated via system form, 2) automated via Hasmukh's bulk skill, or 3) manual",
    },
    {
        "author": "Tony Henriquez",
        "email": "thenriquez@rippling.com",
        "time": "2026-06-09 19:42:44 EDT",
        "text": "Kitty Kwan I am going to run a query in Snowflake to do the deep analysis to determine what needs an amendments vs not, and then break it down in the 3 buckets your outlined. Going to target a first iteration by End of Week if not sooner.",
    },
    {
        "author": "Kitty Kwan",
        "email": "kitty@rippling.com",
        "time": "2026-06-09 19:43:29 EDT",
        "text": "Tony Henriquez thank you - will be helpful to know what type of ops burden we're looking at",
    },
    {
        "author": "Tony Henriquez",
        "email": "thenriquez@rippling.com",
        "time": "2026-06-09 19:44:19 EDT",
        "text": "Hopefully not as bad as we think, but yes for sure.",
        "reactions": "thank_you (1)",
    },
    {
        "author": "Tony Henriquez",
        "email": "thenriquez@rippling.com",
        "time": "2026-06-09 20:12:44 EDT",
        "text": "cc: Robin Islam, Hasmukh Sharma",
        "reactions": "eyes (1)",
    },
    {
        "author": "Tony Henriquez",
        "email": "thenriquez@rippling.com",
        "time": "2026-06-12 17:43:07 EDT",
        "text": "Update: Need to push the MMDD on this to Monday 6/15.",
        "reactions": "+1::skin-tone-2 (1), thank_you (1), pray (1)",
    },
    {
        "author": "Abhishek Shah",
        "email": "abhishah@rippling.com",
        "time": "2026-06-15 00:54:54 EDT",
        "text": """Identify which amendments were user generated vs system generated (DRI: Abhishek Shah)
Nicky DePaul If you're referring to identifying whether amendments were caused by the client or by Rippling for billing purposes, that process is currently handled manually by the Ops team.

At the moment, Engineering does not have a definitive understanding of the criteria used to make that determination. However, if Ops require any necessary data points for that or outline the steps they follow to identify the source of an amendment, I can retrieve the relevant information from our systems. Once the final decision is made, I can also update the `causedBy` field for all affected forms in bulk from the backend.""",
    },
    {
        "author": "Abhishek Shah",
        "email": "abhishah@rippling.com",
        "time": "2026-06-15 00:55:19 EDT",
        "text": """Nicky DePaul Also once the project is shipped, we plan to mark all amendments created from Q3'26 onwards (i.e., createdAt >= 2026-07-01) as eligible for auto-approval.

Important clarification: This is based on the amendment creation date, not the quarter the amendment pertains to.
Example:
• If an amendment form is generated on 2026-07-01 for Q1'25, it will be included in the next auto-approval cycle and will be auto-approved at month-end.

Backlog amendments are not affected by this change.
• Any amendments that were created before 2026-07-01 and remain pending will not be auto-approved.
• In other words, this change is forward-looking only and will not clear or approve existing backlog amendments for a client/agency.

For testing these changes, we'll select a few clients across all of the following categories to ensure we cover all scenarios:
• US and GP clients
• Old-infra and new-infra clients
• Tax-due and refund-due cases
• Automated and manual forms

Rajat Apurwa has already listed down some potential test clients for the above categories and will share the list here shortly.
Please review the list and let me know if you'd like to further filter it based on additional criteria (e.g., ARR, client size, or any other considerations). We can then incorporate the filtered clients into the final testing list that you'll be sharing.

Please let me know if you have questions on any of the above

Cc: Kushal Shah""",
    },
    {
        "author": "Hasmukh Sharma",
        "email": "hsharma@rippling.com",
        "time": "2026-06-15 14:05:07 EDT",
        "text": """Backlog amendments are not affected by this change.
• Any amendments that were created before 2026-07-01 and remain pending will not be auto-approved.
• In other words, this change is forward-looking only and will not clear or approve existing backlog amendments for a client/agency.

Nicky DePaul Tony Henriquez trying to understand if this is what the plan will be!""",
    },
    {
        "author": "Nicky DePaul",
        "email": "ndepaul@rippling.com",
        "time": "2026-06-15 14:30:56 EDT",
        "text": """Abhishek Shah
1. When I say system generated, I'm quite sure there are amendments that are literally generated by the Rippling systems/team, like inc-2025-12-15-incorrect-amendment-form-created-from-runs. Are you saying we cannot identify this?
2. How are we planning to handle the backlog? Will there be a feature flag to enable older amendments to be automatically approved? Resolving the backlog is a primary goal of this project, and we had never aligned that forward-looking was the scope.""",
    },
    {
        "author": "Abhishek Shah",
        "email": "abhishah@rippling.com",
        "time": "2026-06-16 04:36:10 EDT",
        "text": """When I say system generated, I'm quite sure there are amendments that are literally generated by the Rippling systems/team, like inc-2025-12-15-incorrect-amendment-form-created-from-runs. Are you saying we cannot identify this?
Nicky DePaul These are two completely different issues:
1. Root cause of the form generation – whether the form was generated due to a Rippling issue or a client-side action.
2. Form visibility issue – whether a generated form was visible to the client.

The incident you're referring to falls into the second category, where forms had been generated but were not being displayed to the client.
I believe the discussion was more around the first point, specifically to determe which clients can be charged.
At the moment, that determination is made by the Ops team rather than the system. If we can define the rules or criteria that Ops uses for this decision, we can have a first pass at it.


How are we planning to handle the backlog? Will there be a feature flag to enable older amendments to be automatically approved? Resolving the backlog is a primary goal of this project, and we had never aligned that forward-looking was the scope.
Nicky DePaul What I meant is that we can definitely enable auto-approval starting from Q3'26 so that at least any amendments generated from that point onward are automatically approved.
For the backlog of existing amendments, we can incorporate a separate approach once we have a clear strategy in place. Since there is some discussion still ongoing around that, I don't think we should let it block this batch. We can move forward with enabling auto-approval for new amendments while continuing to finalize the backlog handling plan.""",
    },
    {
        "author": "Nicky DePaul",
        "email": "ndepaul@rippling.com",
        "time": "2026-06-16 12:00:24 EDT",
        "text": """Abhishek Shah so zero amendments are generated by "Rippling system"? Like recon runs are generated by Rippling system. Let's confirm this 100%. I understand that we won't be able to tell whether an amendment stemmed from an incident.

Agreed we can start with only go-forward, but just confirming that we can add the full backlog (or any portion we choose) to the batch approval cron whenever we want to?""",
    },
    {
        "author": "Abhishek Shah",
        "email": "abhishah@rippling.com",
        "time": "2026-06-16 12:14:01 EDT",
        "text": """so zero amendments are generated by "Rippling system"? Like recon runs are generated by Rippling system.
Nicky DePaul I believe QWR are Rippling created runs and they do result in amendments. So yes, amendments could be system generated but there are bunch of other factors as well that determine if causedBy should be Client or Rippling. Just the run types may not be sufficient for determining that.

Agreed we can start with only go-forward, but just confirming that we can add the full backlog (or any portion we choose) to the batch approval cron whenever we want to?
Yes, we should be able to do that""",
    },
    {
        "author": "Nicky DePaul",
        "email": "ndepaul@rippling.com",
        "time": "2026-06-16 12:33:54 EDT",
        "text": "QWR is a good example...that would definitely be labeled Rippling right? And thus not charged. We need to gather all of those cases (tbh I think that may be the only one)",
    },
    {
        "author": "James Brown",
        "email": "jabrown@rippling.com",
        "time": "2026-06-18 12:05:47 EDT",
        "text": "Nicky DePaul Lets discuss this on friday. I am concerned about the 7/1 date.",
    },
    {
        "author": "Nicky DePaul",
        "email": "ndepaul@rippling.com",
        "time": "2026-06-18 12:43:27 EDT",
        "text": "we can set whatever date we want right Abhishek Shah?",
    },
    {
        "author": "Abhishek Shah",
        "email": "abhishah@rippling.com",
        "time": "2026-06-18 12:56:21 EDT",
        "text": "we can though that would require one more deployment.\nAny concerns with this date though, just wanted to understand?",
    },
    {
        "author": "Nicky DePaul",
        "email": "ndepaul@rippling.com",
        "time": "2026-06-18 13:03:03 EDT",
        "text": "James Brown can you share more here?",
    },
]


def add_formatted_text(paragraph, text: str) -> None:
    """Add text with basic bold markers (*text*) preserved."""
    parts = re.split(r"(\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_message_block(doc: Document, message: dict, *, is_parent: bool = False) -> None:
    label = "Thread starter" if is_parent else "Reply"
    doc.add_heading(f"{label}: {message['author']}", level=2)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"{message['time']}  |  {message['email']}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for line in message["text"].splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("•") or stripped.startswith("◦"):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, stripped[1:].strip())
        elif re.match(r"^\d+\.", stripped):
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, stripped)
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, stripped)

    if message.get("reactions"):
        reactions = doc.add_paragraph()
        run = reactions.add_run(f"Reactions: {message['reactions']}")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")


def build_document() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_heading(TITLE, level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    info.add_run(f"Channel: {CHANNEL}\n").font.size = Pt(11)
    info.add_run(f"Source: {SLACK_URL}\n").font.size = Pt(10)
    info.add_run(
        f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M %Z')}"
    ).font.size = Pt(10)

    doc.add_paragraph("")
    add_message_block(doc, PARENT, is_parent=True)

    for reply in REPLIES:
        add_message_block(doc, reply)

    return doc


def main() -> Path:
    output = Path("/workspace/output/unsubmitted-amendments-backlog-slack-thread.docx")
    output.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(output)

    artifact = Path("/opt/cursor/artifacts/unsubmitted-amendments-backlog-slack-thread.docx")
    artifact.parent.mkdir(parents=True, exist_ok=True)
    artifact.write_bytes(output.read_bytes())
    return output


if __name__ == "__main__":
    path = main()
    print(f"Created: {path}")
